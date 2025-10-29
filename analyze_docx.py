import os
import markdown
from docx import Document

def analyze_docx_files():
    # Get the current directory
    current_directory = os.getcwd()
    markdown_output = []

    # Loop through all files in the current directory
    for filename in os.listdir(current_directory):
        if filename.endswith('.docx'):
            docx_path = os.path.join(current_directory, filename)
            doc = Document(docx_path)
            analysis = f"# Analysis of {filename}\n\n"

            # Extract text, headings, and tables
            for para in doc.paragraphs:
                analysis += f"{para.text}\n"

            for table in doc.tables:
                for row in table.rows:
                    analysis += '| ' + ' | '.join(cell.text for cell in row.cells) + ' |\n'
                analysis += '\n'

            markdown_output.append(analysis)

    # Write the analysis to markdown files
    for index, output in enumerate(markdown_output):
        output_filename = f"analysis_{index + 1}.md"
        with open(output_filename, 'w') as md_file:
            md_file.write(output)

if __name__ == "__main__":
    analyze_docx_files()